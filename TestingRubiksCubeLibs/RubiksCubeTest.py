#Tests the openai functions on reading from docx files (using docx import), understanding the question asked, and returning a correct answer with references/sources...
#also tests converting .pdf files to .docx files...

from urllib import response
from openai import OpenAI   #imports openai...
from docx import Document
import docx     #to read docx documents...*Note: attempts to read from pdf docs results in errors...docx files seem easier to read for python...
from docx.table import _Cell, Table     #opportunity to use specific cell/table functions for higher usage of docx import...
from pdf2docx import Converter      #to convert from .pdf to .docx file...
from docxcompose.composer import Composer
from docx import Document as Document_compose
from pathlib import Path


import os



ROOT_DIR = Path(__file__).parent
document7=ROOT_DIR/'Cube_Notation_Docx.docx'
document8=ROOT_DIR/'Solving_Cube_Beginners_Method_Docx.docx'
document9=ROOT_DIR/'Step1_Docx.docx'

client = OpenAI()   #assign client for use for openai...


def getText():     #reads text from normal paragraph format...
    doc = docx.Document(document8)  #loads the document to use...
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def getText2():     #reads text from normal paragraph format...
    doc = docx.Document(document7)  #loads the document to use...
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def getText3():     #reads text from normal paragraph format...
    doc = docx.Document(document9)  #loads the document to use...
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def getTableText():     #reads text from tables...
    doc = docx.Document(document8)  #loads the document to use...
    fullText = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                fullText.append(cell.text)
                #print(cell.text)
    return '\n'.join(fullText)

def getTableText2():     #reads text from tables...
    doc = docx.Document(document7)  #loads the document to use...
    fullText = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                fullText.append(cell.text)
                #print(cell.text)
    return '\n'.join(fullText)

def getTableText3():     #reads text from tables...
    doc = docx.Document(document9)  #loads the document to use...
    fullText = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                fullText.append(cell.text)
                #print(cell.text)
    return '\n'.join(fullText)


# query3 = f"""Solve the Rubiks cube from the cube notation that the user provides.  In order to solve the cube, you must perform the same move over and over again until the cube is fully solved and the cube configuration turns into 
#     YYYYYYYYYRRRRRRRRRGGGGGGGGGOOOOOOOOOBBBBBBBBBWWWWWWWWW.  Unless the cube is in this state, then keep performing the same move again and again until it reaches this state.  The move you must perform over and over again is F' L' F' L and you must only
#     perform this one move multiple times until the cube notation shows it is solved.  After each move you perform, output the current cube notation.  Repeat this until the cube is solved.  

#     Files:
#     \"\"\"
#     {getTableText(), getText()}     
#     \"\"\"

#     Question: {question}"""

# import magiccube

# # 3x3x3 Cube
# cube = magiccube.Cube(3,"YYYYYYYYYRRRRRRRRRGGGGGGGGGOOOOOOOOOBBBBBBBBBWWWWWWWWW")
# print(cube) OYYYYYYYYRRRRRRRRRGGGGGGGGGOOOOOOOOOBBBBBBBBBWWWWWWWWW
# OGOBBYWROGOBYOWWGBOYBBWGYOWRWROGBROYYBGRRRBRGWWYWYGRYG

question=input("What question do you want to ask the chatbot?\n")   #asks user question 
#=For each sentence in your reasoning, compare it to the previous sentence.  If there is a discontinuity in reasoning for the sentences, only print the new sentence and discard the previous one.
#If the questions asks for an answer regarding the class meeting days or times, make sure to check if the file mentions the class being online or in-person.
def queryFinal():
    query3 = f"""Solve the Rubiks cube from the cube notation that the user provides.

    Files:
    \"\"\"
    {getText3(), getText2(), getText(), getTableText(), getTableText2(), getTableText3()}     
    \"\"\"

    Question: {question}"""     #this version 'query3' works best when taking text from .docx files with tables contained...'getTableText' helps, but query3 contains specific wording on displaying the information the bot retrieves...

    response = client.chat.completions.create(
    messages=[
        {'role': 'system', 'content': 'You solve a Rubiks cube using by only performing the same move over and over again.'},
        {'role': 'user', 'content': query3},    #uses 'query3' for content to respond with...
    ],
    model="gpt-3.5-turbo",  #model to use...
    temperature=0,
    )
    answer = response.choices[0].message.content
    print(answer)      #answers user's question...

    
while question != 'exit':   #loops so user can ask different questions...
    queryFinal()    #if != 'exit' then will go through chatbot...
    question=""
    question=input("What question do you want to ask the chatbot?\n")   #asks user again for question...